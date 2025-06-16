from github import Github
from datetime import datetime
import config
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import pdfkit
import os
from collections import defaultdict
import plotly.express as px
from ai_helpers import generate_ai_description
import markdown2
import warnings

# Suppress warnings
warnings.filterwarnings("ignore")

print("Script started!")
print("🟢 Phase 1: Config check")
print(f"GitHub user: {config.GITHUB_USERNAME}")
print(f"Token exists: {bool(config.GITHUB_TOKEN)}")

print("🟢 Phase 2: GitHub connection")
g = Github(config.GITHUB_TOKEN, timeout=30)
user = g.get_user(config.GITHUB_USERNAME)
print(f"Connected to: {user.login}")
print(f"Public repos: {user.public_repos}")

print("🟢 Phase 3: Getting repos")
repos = list(user.get_repos())
print(f"Found {len(repos)} repositories. First repo: {repos[0].name if repos else 'None'}")

class RepoDocumentGenerator:
    def __init__(self):
        self.g = Github(config.GITHUB_TOKEN, timeout=30)
        self.user = self.g.get_user(config.GITHUB_USERNAME)
        
    def get_repos(self, filters=None, sort_by='updated'):
        """Get repositories with optional filtering and sorting"""
        repos = list(self.user.get_repos())
        
        if filters:
            if filters.get('language'):
                repos = [r for r in repos if r.language == filters['language']]
            if filters.get('min_stars'):
                repos = [r for r in repos if r.stargazers_count >= filters['min_stars']]
        
        if sort_by == 'stars':
            repos.sort(key=lambda x: x.stargazers_count, reverse=True)
        elif sort_by == 'forks':
            repos.sort(key=lambda x: x.forks_count, reverse=True)
        else:
            repos.sort(key=lambda x: x.updated_at, reverse=True)
            
        return repos
    
    def get_repo_metrics(self, repo):
        """Get extended metrics with AI-generated description"""
        metrics = {
            'name': repo.name,
            'description': repo.description,
            'url': repo.html_url,
            'stars': repo.stargazers_count,
            'forks': repo.forks_count,
            'language': repo.language,
            'updated_at': repo.updated_at,
            'topics': repo.get_topics(),
            'has_issues': repo.has_issues,
            'open_issues': repo.open_issues_count,
            'has_wiki': repo.has_wiki,
            'archived': repo.archived,
            'license': repo.license.name if repo.license else None,
            'contributors': len(list(repo.get_contributors())),
            'pull_requests': len(list(repo.get_pulls(state='all'))),
            'size': repo.size,
            'ai_description': None
        }
        
        try:
            readme = repo.get_readme()
            readme_content = readme.decoded_content.decode('utf-8')
            metrics['readme'] = readme_content
            metrics['ai_description'] = generate_ai_description(repo.name, readme_content)
        except Exception as e:
            metrics['readme'] = None
            print(f"⚠️ Failed to process {repo.name}: {str(e)}")
            
        return metrics
    
    def generate_language_chart(self, repos, output_path):
        """Generate language distribution chart"""
        languages = [r.language for r in repos if r.language]
        lang_counts = defaultdict(int)
        for lang in languages:
            lang_counts[lang] += 1
            
        plt.figure(figsize=(10, 6))
        sns.barplot(x=list(lang_counts.values()), y=list(lang_counts.keys()))
        plt.title('Language Distribution Across Repositories')
        plt.xlabel('Number of Repositories')
        plt.tight_layout()
        plt.savefig(f'{output_path}/language_distribution.png')
        plt.close()
        
        fig = px.pie(names=list(lang_counts.keys()), 
                    values=list(lang_counts.values()),
                    title='Language Distribution')
        fig.write_html(f'{output_path}/language_pie.html')
    
    def generate_activity_chart(self, repos, output_path):
        """Generate activity timeline chart"""
        dates = [r.updated_at.date() for r in repos]
        counts = defaultdict(int)
        for date in dates:
            counts[date] += 1
            
        dates_sorted = sorted(counts.keys())
        values = [counts[d] for d in dates_sorted]
        
        plt.figure(figsize=(12, 6))
        plt.plot(dates_sorted, values)
        plt.title('Repository Update Activity')
        plt.xlabel('Date')
        plt.ylabel('Number of Updates')
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig(f'{output_path}/activity_timeline.png')
        plt.close()
    
    def generate_markdown(self, repos_metrics):
        """Generate markdown with enhanced descriptions"""
        doc = "# GitHub Repository Summary\n\n"
        doc += f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
        
        # Summary statistics
        doc += "## 📊 Summary Statistics\n"
        doc += f"- 📦 Total Repositories: {len(repos_metrics)}\n"
        doc += f"- ⭐ Total Stars: {sum(r['stars'] for r in repos_metrics)}\n"
        doc += f"- 🍴 Total Forks: {sum(r['forks'] for r in repos_metrics)}\n"
        doc += f"- 👥 Total Contributors: {sum(r['contributors'] for r in repos_metrics)}\n"
        doc += f"- 🔀 Total Pull Requests: {sum(r['pull_requests'] for r in repos_metrics)}\n\n"
        
        # Repository details
        doc += "## 📚 Repository Details\n\n"
        for repo in repos_metrics:
            doc += f"### 🗂️ [{repo['name']}]({repo['url']})\n\n"
            
            # Description section
            doc += "#### 📝 Description\n"
            if repo['description']:
                doc += f"- **Original Description:** {repo['description']}\n"
            if repo['ai_description']:
                doc += f"- **AI-Generated Description:** {repo['ai_description']}\n"
            if not repo['description'] and not repo['ai_description']:
                doc += "- No description available\n"
            doc += "\n"
            
            # Statistics section
            doc += "#### 📈 Statistics\n"
            doc += f"- ⭐ Stars: {repo['stars']}\n"
            doc += f"- 🍴 Forks: {repo['forks']}\n"
            doc += f"- 📝 Language: {repo['language'] or 'Not specified'}\n"
            doc += f"- 📅 Last Updated: {repo['updated_at'].strftime('%Y-%m-%d')}\n"
            doc += f"- 🏷️ Topics: {', '.join(repo['topics']) if repo['topics'] else 'None'}\n"
            doc += f"- 📜 License: {repo['license'] or 'None'}\n"
            doc += f"- 👥 Contributors: {repo['contributors']}\n"
            doc += f"- 🔀 Pull Requests: {repo['pull_requests']}\n"
            doc += f"- 🐛 Open Issues: {repo['open_issues']}\n"
            doc += f"- 📦 Size: {round(repo['size']/1024, 2)} MB\n\n"
            
            # README preview
            if repo['readme']:
                doc += "#### 📖 README Preview\n```\n"
                doc += "\n".join(repo['readme'].split("\n")[:5])  # First 5 lines
                doc += "\n...\n```\n\n"
            
            doc += "---\n\n"
        
        return doc
    
    def generate_html_report(self, repos_metrics, output_path):
        """Generate HTML report with embedded charts"""
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>GitHub Repositories Report</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                .repo {{ margin-bottom: 2rem; border-bottom: 1px solid #eee; padding-bottom: 1rem; }}
                .chart {{ margin: 2rem 0; }}
                .ai-description {{ background-color: #f8f9fa; padding: 1rem; border-radius: 5px; }}
            </style>
        </head>
        <body>
            <h1>GitHub Repository Summary</h1>
            <p>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            
            <div class="chart">
                <h2>Language Distribution</h2>
                <img src="language_distribution.png" alt="Language Distribution" style="max-width: 100%;">
            </div>
            
            <div class="chart">
                <h2>Activity Timeline</h2>
                <img src="activity_timeline.png" alt="Activity Timeline" style="max-width: 100%;">
            </div>
            
            {markdown2.markdown(self.generate_markdown(repos_metrics))}
        </body>
        </html>
        """
        
        with open(f'{output_path}/repos_summary.html', 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def generate_pdf(self, markdown_content, output_path):
        """Generate PDF using path from .env"""
        if not hasattr(config, 'WKHTMLTOPDF_PATH'):
            print("⚠️ WKHTMLTOPDF_PATH not configured in .env")
            return False
            
        wkhtml_path = config.WKHTMLTOPDF_PATH
        
        if not os.path.exists(wkhtml_path):
            print(f"❌ wkhtmltopdf.exe not found at: {wkhtml_path}")
            print("Please verify the path in your .env file")
            return False
            
        try:
            pdf_config = pdfkit.configuration(wkhtmltopdf=wkhtml_path)
            html = markdown2.markdown(markdown_content)
            pdfkit.from_string(html, output_path, configuration=pdf_config)
            return True
        except Exception as e:
            print(f"⚠️ PDF generation failed: {str(e)}")
            return False
    
    def generate_docx(self, markdown_content, output_path):
        """Generate Word document"""
        doc = Document()
        doc.add_heading('GitHub Repository Summary', 0)
        
        for line in markdown_content.split('\n'):
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=3)
            elif line.startswith('- **'):
                # Handle bold list items
                text = line[2:].replace('**', '')
                doc.add_paragraph(text, style='List Bullet')
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('```'):
                continue  # Skip code blocks in Word doc
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
    
    def generate_all(self, output_dir='output', filters=None, sort_by='updated'):
        """Generate all report formats"""
        os.makedirs(output_dir, exist_ok=True)
        
        repos = self.get_repos(filters, sort_by)
        repos_metrics = [self.get_repo_metrics(r) for r in repos]
        
        self.generate_language_chart(repos, output_dir)
        self.generate_activity_chart(repos, output_dir)
        
        markdown_content = self.generate_markdown(repos_metrics)
        
        with open(f'{output_dir}/repos_summary.md', 'w', encoding='utf-8') as f:
            f.write(markdown_content)
            
        self.generate_pdf(markdown_content, f'{output_dir}/repos_summary.pdf')
        self.generate_docx(markdown_content, f'{output_dir}/repos_summary.docx')
        self.generate_html_report(repos_metrics, output_dir)
        
        return {
            'markdown': f'{output_dir}/repos_summary.md',
            'pdf': f'{output_dir}/repos_summary.pdf',
            'docx': f'{output_dir}/repos_summary.docx',
            'html': f'{output_dir}/repos_summary.html',
            'charts': {
                'language': f'{output_dir}/language_distribution.png',
                'activity': f'{output_dir}/activity_timeline.png'
            }
        }

if __name__ == "__main__":
    try:
        print("\n🔧 Creating generator instance...")
        generator = RepoDocumentGenerator()

        print("\n🔍 Testing Document Generation")
        if repos:
            print("🟢 Testing markdown generation...")
            test_repo = generator.get_repo_metrics(repos[0])
            test_md = generator.generate_markdown([test_repo])
            print(f"Markdown sample:\n{test_md[:300]}...")
            
            print("\n🔍 AI Description Test:")
            if test_repo['ai_description']:
                print(test_repo['ai_description'])
            else:
                print("⚠️ No AI description generated")
            
            print("\n🟢 Testing chart generation...")
            generator.generate_language_chart(repos[:10], "output")
            print("Chart generated!")
            
            print("🟢 Testing PDF generation...")
            if generator.generate_pdf("# Test\n\nAI Description Test", "output/test.pdf"):
                print("PDF generated successfully!")
            else:
                print("⚠️ PDF generation failed or not available")
            
            print("\n🟢 Running full generation...")
            results = generator.generate_all()
            if results:
                print(f"Successfully generated documents in {results['markdown']}")
            else:
                print("❌ Document generation failed")
        else:
            print("❌ No repositories available for testing")
            
    except Exception as e:
        print(f"❌ Fatal error: {str(e)}")
        import traceback
        traceback.print_exc()